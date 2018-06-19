from django.shortcuts import render
from django.http import HttpResponse
from django.shortcuts import redirect
from django.urls import reverse
import subprocess
import docx
import os

from zhilianjob.models import goods
from zhilianjob.models import PosiN
from zhilianjob.models import CompN
from zhilianjob.models import WorkEduN
from zhilianjob.models import WorkN
from zhilianjob.models import SizeN
from zhilianjob.models import AttributeN
from zhilianjob.models import FieldN
from zhilianjob.models import EduSalN
from zhilianjob.models import ExpSalN
from zhilianjob.models import NewJobs

from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LTTextBoxHorizontal, LAParams
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed
from familia_wrapper import InferenceEngineWrapper


# Create your views here.


def test(request):
    # string = "因为我是天才啊"
    # TutorialList = ["HTML", "CSS", "JQuery", "Python", "Giango"]
    # info_dict = {'site': '第一个字典的值', 'content': '第二个字典的值'}

    # List = map(str, range(100))
    # return render(request, 'test.html', {'List': List})
    goods_list = list(goods.objects.all())
    print(goods_list)
    return render(request, "zhu.html", {"goods_list": goods_list})


def index(request):
    return render(request, "index.html")


def page_not_found(request):
    return render(request, "404.html")


def posin(request):
    PosiN_list = list(PosiN.objects.all().order_by('-num'))
    return render(request, "echarts/posin.html", {"posin_list": PosiN_list})


def workn(request):
    WorkN_list = list(WorkN.objects.all().order_by('-num'))
    return render(request, "echarts/workn.html", {"workn_list": WorkN_list})


def compn(request):
    CompN_list = list(CompN.objects.all().order_by('-num'))
    return render(request, "echarts/compn.html", {"workn_list": CompN_list})


def sizen(request):
    SizeN_list = list(SizeN.objects.all().order_by('-num'))
    return render(request, "echarts/sizen.html", {"sizen_list": SizeN_list})


def attributen(request):
    AttributeN_list = list(AttributeN.objects.all())
    return render(request, "echarts/attributen.html", {"attributen_list": AttributeN_list})


def fieldn(request):
    FieldN_list = list(FieldN.objects.all())
    return render(request, "echarts/fieldn.html", {"fieldn_list": FieldN_list})


def workedun(request):
    list_buxian = list(WorkEduN.objects.filter(education='不限').order_by('-workplace'))
    list_zhongzhuan = list(WorkEduN.objects.filter(education='中专').order_by('-workplace'))
    list_zhongji = list(WorkEduN.objects.filter(education='中技').order_by('-workplace'))
    list_qita = list(WorkEduN.objects.filter(education='其他').order_by('-workplace'))
    list_boshi = list(WorkEduN.objects.filter(education='博士').order_by('-workplace'))
    list_dazhuan = list(WorkEduN.objects.filter(education='大专').order_by('-workplace'))
    list_benke = list(WorkEduN.objects.filter(education='本科').order_by('-workplace'))
    list_shuoshi = list(WorkEduN.objects.filter(education='硕士').order_by('-workplace'))
    list_gaozhong = list(WorkEduN.objects.filter(education='高中').order_by('-workplace'))
    return render(request, "echarts/workedun.html",
                  {"list_buxian": list_buxian, "list_zhongzhuan": list_zhongzhuan, "list_zhongji": list_zhongji,
                   "list_qita": list_qita, "list_boshi": list_boshi, "list_dazhuan": list_dazhuan,
                   "list_benke": list_benke, "list_shuoshi": list_shuoshi, "list_gaozhong": list_gaozhong})


def edusaln(request):
    list_8000_10000 = list(EduSalN.objects.filter(salary='8001-10000元/月').order_by('education'))
    list_6000_8000 = list(EduSalN.objects.filter(salary='6001-8000元/月').order_by('education'))
    list_4000_6000 = list(EduSalN.objects.filter(salary='4001-6000元/月').order_by('education'))
    list_10000_15000 = list(EduSalN.objects.filter(salary='10001-15000元/月').order_by('education'))
    list_2000_4000 = list(EduSalN.objects.filter(salary='2001-4000元/月').order_by('education'))
    list_face = list(EduSalN.objects.filter(salary='面议').order_by('education'))
    list_10000_20000 = list(EduSalN.objects.filter(salary='10000-20000元/月').order_by('education'))
    list_1000 = list(EduSalN.objects.filter(salary='1000-1001元/月').order_by('education'))
    list_3300_4500 = list(EduSalN.objects.filter(salary='3300-4500元/月').order_by('education'))
    list_other = list(EduSalN.objects.filter(salary='其他').order_by('education'))
    return render(request, "echarts/edusaln.html",
                  {"list_8000_10000": list_8000_10000, "list_6000_8000": list_6000_8000,
                   "list_4000_6000": list_4000_6000,
                   "list_10000_15000": list_10000_15000, "list_2000_4000": list_2000_4000, "list_face": list_face,
                   "list_10000_20000": list_10000_20000, "list_1000": list_1000, "list_3300_4500": list_3300_4500,
                   "list_other": list_other})


def expsaln(request):
    list_6000_8000 = list(ExpSalN.objects.filter(salary='6001-8000元/月').order_by('experience'))
    list_2000_4000 = list(ExpSalN.objects.filter(salary='2001-4000元/月').order_by('experience'))
    list_4000_6000 = list(ExpSalN.objects.filter(salary='4001-6000元/月').order_by('experience'))
    list_2600_2600 = list(ExpSalN.objects.filter(salary='2600-2600元/月').order_by('experience'))
    list_3000_6000 = list(ExpSalN.objects.filter(salary='3000-6000元/月').order_by('experience'))
    list_other = list(ExpSalN.objects.filter(salary='其他').order_by('experience'))
    list_8000_10000 = list(ExpSalN.objects.filter(salary='8001-10000元/月').order_by('experience'))
    list_10000_15000 = list(ExpSalN.objects.filter(salary='10001-15000元/月').order_by('experience'))
    list_15000_20000 = list(ExpSalN.objects.filter(salary='15001-20000元/月').order_by('experience'))
    return render(request, "echarts/expsaln.html", {"list_6000_8000": list_6000_8000, "list_2000_4000": list_2000_4000,
                                                    "list_4000_6000": list_4000_6000, "list_2600_2600": list_2600_2600,
                                                    "list_3000_6000": list_3000_6000, "list_other": list_other,
                                                    "list_8000_10000": list_8000_10000,
                                                    "list_10000_15000": list_10000_15000,
                                                    "list_15000_20000": list_15000_20000})


# def hello(request):
#     return HttpResponse('我喜欢你了')


def recommend_cal(text):
    short_engine_wrapper = InferenceEngineWrapper('/root/Familia/model/webpage', 'lda.conf', 'webpage_twe_lda.model')
    doc_seg_short = short_engine_wrapper.tokenize(text)

    long_engine_wrapper = InferenceEngineWrapper('/root/Familia/model/webpage', 'lda.conf')
    doc_seg_long = long_engine_wrapper.tokenize(text)

    top_200_list = short_long_cal(short_engine_wrapper, doc_seg_short)

    top_3_jobs = long_long_cal(long_engine_wrapper, doc_seg_long, top_200_list)
    return top_3_jobs


def short_long_cal(short_engine_wrapper, doc_seg_short):
    list_new_jobs = list(NewJobs.objects.all()[:6000])
    temp_dict = {}
    index = 0
    for jobs in list_new_jobs:
        query_seg = short_engine_wrapper.tokenize(jobs.position)
        twe_similarity = short_engine_wrapper.cal_query_doc_similarity(query_seg, doc_seg_short)
        temp_dict[jobs] = twe_similarity[1]
        print(str(index) + "\t" + str(twe_similarity[1]))
        index += 1
    print(''''
    ***************************************************************************************************
    ***************************************************************************************************
    ***************************************************************************************************
    ***************************************************************************************************
    ***************************************************************************************************
    ''')
    top_200_list = sorted(temp_dict.items(), key=lambda e: e[1], reverse=True)[:200]
    return top_200_list


def long_long_cal(long_engine_wrapper, doc_seg_long, top_200_list):
    temp_dict = {}
    index = 0
    for item in top_200_list:
        doc_seg_description = long_engine_wrapper.tokenize(item[0].description)
        distance = long_engine_wrapper.cal_doc_distance(doc_seg_long, doc_seg_description)
        # temp_dict[item[0]] = distance[0]                    # jensen_shannon_divergence
        temp_dict[item[0]] = distance[1]  # hellinger_distance
        print(str(index) + "\t" + str(distance[1]))
        index += 1
    top_3_list = sorted(temp_dict.items(), key=lambda e: e[1], reverse=False)[:3]
    top_3_jobs = []
    for jobs in top_3_list:
        top_3_jobs.append(jobs[0])
    return top_3_jobs


def upload(request):
    if request.method == "POST":
        try:
            file_path = save_upload_file(request.FILES['file'], str(request.FILES['file']))
            flag, text = get_file_content(file_path)
            if flag == 0:
                request.session['text'] = text
                return redirect(reverse('result'))
                # return HttpResponse('Succeed ' + text)  # 此处简单返回一个成功的消息，在实际应用中可以返回到指定的页面中
            else:
                return render(request, "404.html", {"text": 'Failed ' + text + '\n'})
                # return HttpResponse('Failed ' + text)  # 此处简单返回一个简历格式错误的消息，在实际应用中可以返回到指定的页面中
        except Exception as e:
            text = '您还没有上传简历，或者上传错误，请重新上传简历'
            return render(request, "404.html", {"text": 'Failed ' + text + '\n' + str(e)})
            # return HttpResponse('Failed ' + text + '\n' + str(e))  # 此处简单返回一个简历上传异常的消息，在实际应用中可以返回到指定的页面中

    return render(request, "upload_file.html")


def result(request):
    top_3_jobs = recommend_cal(request.session['text'])
    print(top_3_jobs)
    return render(request, "result.html", {"top_3_jobs": top_3_jobs})


def save_upload_file(file, filename):
    path = 'media/uploads/'  # 上传文件的保存路径，可以自己指定任意的路径
    if not os.path.exists(path):
        os.makedirs(path)
    with open(path + filename, 'wb+') as destination:
        for chunk in file.chunks():
            destination.write(chunk)
        file_path = destination.name
    return file_path


def get_file_content(file_path):
    text = ''
    # flag = -1   # -1失败，0成功
    if str(file_path).endswith(".docx"):
        doc = docx.Document(file_path)
        for graph in doc.paragraphs:
            text += graph.text
        flag = 0
    elif str(file_path).endswith('.pdf'):
        fp = open(file_path, 'rb')  # 以二进制读模式打开
        praser = PDFParser(fp)  # 用文件对象来创建一个pdf文档分析器
        doc = PDFDocument()  # 创建一个PDF文档
        # 连接分析器 与文档对象
        praser.set_document(doc)
        doc.set_parser(praser)

        doc.initialize()  # 提供初始化密码 如果没有密码 就创建一个空的字符串

        # 检测文档是否提供txt转换，不提供就忽略
        if not doc.is_extractable:
            raise PDFTextExtractionNotAllowed
        else:
            rsrcmgr = PDFResourceManager()  # 创建PDf 资源管理器 来管理共享资源

            # 创建一个PDF设备对象
            laparams = LAParams()
            device = PDFPageAggregator(rsrcmgr, laparams=laparams)
            interpreter = PDFPageInterpreter(rsrcmgr, device)  # 创建一个PDF解释器对象

            # 循环遍历列表，每次处理一个page的内容
            for page in doc.get_pages():  # doc.get_pages() 获取page列表
                interpreter.process_page(page)
                # 接受该页面的LTPage对象
                layout = device.get_result()
                # 这里layout是一个LTPage对象 里面存放着 这个page解析出的各种对象 一般包括LTTextBox, LTFigure, LTImage, LTTextBoxHorizontal
                # 等等 想要获取文本就获得对象的text属性，
                for x in layout:
                    if isinstance(x, LTTextBoxHorizontal):
                        text += x.get_text()
            fp.close()
        flag = 0
    elif str(file_path).endswith('.doc'):
        text = subprocess.check_output(['antiword', file_path]).decode()
        print(text)
        flag = 0
    else:
        text = '您的简历格式错误，请重新上传简历，支持docx,doc,pdf格式'
        flag = -1

    os.remove(file_path)
    return flag, text
